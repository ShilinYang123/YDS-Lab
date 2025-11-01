#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
批量DOCX文档内容替换工具
用于批量替换指定目录下所有docx文件中的特定文本内容
"""

import os
import sys
from pathlib import Path
from docx import Document
import logging
from typing import Dict, List, Tuple

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('batch_docx_replacer.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

class BatchDocxReplacer:
    """批量DOCX文档替换器"""
    
    def __init__(self, source_dir: str):
        self.source_dir = Path(source_dir)
        self.replacement_map = {
            '生产部': '装配部',
            '市场部': '业务部', 
            '货仓': '仓库',
            '质安部': '品质部',
            '制造部': '五金、装配部',
            '车间组': '生产线',
            '班组': '拉线',
            '设计部': '研发部',
            '供销科': '业务部',
            '检验科': '品质部',
            '车间组长': '拉长',
            '修改员': '维修员',
            '调机员': '夹具设备维护员',
            '制模工': '夹具设备维护员',
            '领班': '拉长',
            '班组长': '拉长',
            '车间负责人': '生产主管',
            '仓务员': '仓管员',
            '付总': '副总'
        }
        self.processed_files = []
        self.error_files = []
        
    def find_docx_files(self) -> List[Path]:
        """递归查找所有docx文件"""
        docx_files = []
        for root, dirs, files in os.walk(self.source_dir):
            for file in files:
                if file.endswith('.docx') and not file.startswith('~$'):
                    docx_files.append(Path(root) / file)
        return docx_files
    
    def replace_text_in_paragraph(self, paragraph, replacements: Dict[str, str]) -> int:
        """在段落中替换文本"""
        replace_count = 0
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                # 处理段落中的runs
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        replace_count += 1
        return replace_count
    
    def replace_text_in_table(self, table, replacements: Dict[str, str]) -> int:
        """在表格中替换文本"""
        replace_count = 0
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_count += self.replace_text_in_paragraph(paragraph, replacements)
        return replace_count
    
    def process_document(self, file_path: Path) -> Tuple[bool, int]:
        """处理单个文档"""
        try:
            logger.info(f"正在处理文件: {file_path}")
            doc = Document(file_path)
            total_replacements = 0
            
            # 处理段落
            for paragraph in doc.paragraphs:
                total_replacements += self.replace_text_in_paragraph(paragraph, self.replacement_map)
            
            # 处理表格
            for table in doc.tables:
                total_replacements += self.replace_text_in_table(table, self.replacement_map)
            
            # 保存文档
            doc.save(file_path)
            
            if total_replacements > 0:
                logger.info(f"文件 {file_path.name} 完成替换，共替换 {total_replacements} 处")
            else:
                logger.info(f"文件 {file_path.name} 无需替换")
                
            return True, total_replacements
            
        except Exception as e:
            logger.error(f"处理文件 {file_path} 时出错: {str(e)}")
            return False, 0
    
    def process_all_documents(self) -> Dict[str, any]:
        """批量处理所有文档"""
        logger.info(f"开始批量处理目录: {self.source_dir}")
        
        docx_files = self.find_docx_files()
        logger.info(f"找到 {len(docx_files)} 个docx文件")
        
        total_files = len(docx_files)
        successful_files = 0
        total_replacements = 0
        
        for file_path in docx_files:
            success, replacements = self.process_document(file_path)
            if success:
                successful_files += 1
                total_replacements += replacements
                self.processed_files.append(str(file_path))
            else:
                self.error_files.append(str(file_path))
        
        # 生成处理报告
        report = {
            'total_files': total_files,
            'successful_files': successful_files,
            'failed_files': len(self.error_files),
            'total_replacements': total_replacements,
            'processed_files': self.processed_files,
            'error_files': self.error_files,
            'replacement_map': self.replacement_map
        }
        
        logger.info(f"批量处理完成: 总文件数={total_files}, 成功={successful_files}, 失败={len(self.error_files)}, 总替换数={total_replacements}")
        
        return report
    
    def generate_report(self, report: Dict[str, any], output_path: str = None):
        """生成处理报告"""
        if output_path is None:
            output_path = self.source_dir.parent / "批量替换处理报告.md"
        
        report_content = f"""# 批量DOCX文档替换处理报告

## 处理概况
- 处理目录: {self.source_dir}
- 总文件数: {report['total_files']}
- 成功处理: {report['successful_files']}
- 处理失败: {report['failed_files']}
- 总替换次数: {report['total_replacements']}

## 替换规则
"""
        
        for old_text, new_text in report['replacement_map'].items():
            report_content += f"- {old_text} → {new_text}\n"
        
        if report['processed_files']:
            report_content += "\n## 成功处理的文件\n"
            for file_path in report['processed_files']:
                report_content += f"- {Path(file_path).name}\n"
        
        if report['error_files']:
            report_content += "\n## 处理失败的文件\n"
            for file_path in report['error_files']:
                report_content += f"- {Path(file_path).name}\n"
        
        report_content += f"\n## 处理时间\n{logging.Formatter().formatTime(logging.LogRecord('', 0, '', 0, '', (), None))}\n"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(report_content)
        
        logger.info(f"处理报告已生成: {output_path}")

def main():
    """主函数"""
    # 默认处理目录，可根据YDS-Lab项目需求调整
    source_directory = "S:\\YDS-Lab\\Docs"
    
    # 创建替换器实例
    replacer = BatchDocxReplacer(source_directory)
    
    # 执行批量替换
    report = replacer.process_all_documents()
    
    # 生成报告
    replacer.generate_report(report)
    
    print(f"\n批量替换完成!")
    print(f"总文件数: {report['total_files']}")
    print(f"成功处理: {report['successful_files']}")
    print(f"处理失败: {report['failed_files']}")
    print(f"总替换次数: {report['total_replacements']}")

if __name__ == "__main__":
    main()