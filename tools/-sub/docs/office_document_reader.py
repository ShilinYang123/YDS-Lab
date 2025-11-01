#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
YDS-Lab项目 - Office文档读取工具
专门用于读取Excel、Word、PowerPoint等Office文档
杨老师要求固化的MCP工具替代方案
"""

import os
import sys
import pandas as pd
from pathlib import Path
try:
    import openpyxl
    from docx import Document
    import antiword
except ImportError:
    print("正在安装必要的依赖包...")
    os.system("pip install openpyxl python-docx pandas antiword")
    import openpyxl
    from docx import Document
    import antiword

class OfficeDocumentReader:
    """Office文档读取器"""
    
    def __init__(self):
        self.supported_formats = ['.xlsx', '.xls', '.docx', '.doc']
    
    def read_excel(self, file_path):
        """读取Excel文件"""
        try:
            # 读取所有工作表
            excel_file = pd.ExcelFile(file_path)
            result = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                result[sheet_name] = df
                
            return result
        except Exception as e:
            print(f"读取Excel文件失败: {e}")
            return None
    
    def read_word(self, file_path):
        """读取Word文档"""
        file_path = Path(file_path)
        
        # 处理.doc文件
        if file_path.suffix.lower() == '.doc':
            try:
                # 使用antiword读取.doc文件
                text = antiword.extract(str(file_path))
                if text:
                    # 按行分割文本
                    paragraphs = [line.strip() for line in text.split('\n') if line.strip()]
                    return {
                        'paragraphs': paragraphs,
                        'tables': []  # .doc文件暂不支持表格提取
                    }
            except Exception as e:
                print(f"使用antiword读取.doc文件失败: {e}")
        
        # 处理.docx文件
        try:
            doc = Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            # 读取表格
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                'paragraphs': content,
                'tables': tables
            }
        except Exception as e:
            print(f"读取Word文档失败: {e}")
            return None
    
    def read_document(self, file_path):
        """统一文档读取接口"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            print(f"文件不存在: {file_path}")
            return None
        
        suffix = file_path.suffix.lower()
        
        if suffix in ['.xlsx', '.xls']:
            return self.read_excel(file_path)
        elif suffix in ['.docx', '.doc']:
            return self.read_word(file_path)
        else:
            print(f"不支持的文件格式: {suffix}")
            return None
    
    def format_output(self, data, file_path):
        """格式化输出"""
        print(f"\n{'='*80}")
        print(f"文件: {file_path}")
        print(f"{'='*80}")
        
        if isinstance(data, dict):
            if 'paragraphs' in data:  # Word文档
                print("\n【文档内容】")
                for i, para in enumerate(data['paragraphs'], 1):
                    print(f"{i:3d}. {para}")
                
                if data['tables']:
                    print("\n【表格内容】")
                    for i, table in enumerate(data['tables'], 1):
                        print(f"\n表格 {i}:")
                        for row in table:
                            print("  | " + " | ".join(row) + " |")
            
            else:  # Excel文件
                for sheet_name, df in data.items():
                    print(f"\n【工作表: {sheet_name}】")
                    if not df.empty:
                        print(df.to_string(index=False))
                    else:
                        print("(空工作表)")
        
        print(f"\n{'='*80}")
        print("文档读取完成")
        print(f"{'='*80}\n")

def main():
    """主函数"""
    if len(sys.argv) != 2:
        print("使用方法: python office_document_reader.py <文件路径>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    reader = OfficeDocumentReader()
    
    data = reader.read_document(file_path)
    if data:
        reader.format_output(data, file_path)
    else:
        print("文档读取失败")

if __name__ == "__main__":
    main()